import os
import re
import tempfile
import traceback
from pathlib import Path
from typing import Dict, List

import numpy as np
from flask import Flask, flash, render_template, request
from langchain import LLMChain, PromptTemplate
from langchain.chains.summarize import load_summarize_chain
from langchain.schema import Document
from langchain_community.document_loaders import (
    PyMuPDFLoader,
    Docx2txtLoader,
)
from langchain_community.embeddings import CohereEmbeddings
from langchain_community.llms import Cohere


COHERE_API_KEY = os.getenv("COHERE_API_KEY", "HIDDEN")
if not COHERE_API_KEY:
    raise RuntimeError("COHERE_API_KEY env-var missing")

llm = Cohere(cohere_api_key=COHERE_API_KEY)
embedder = CohereEmbeddings(cohere_api_key=COHERE_API_KEY, user_agent="quizly-app")


mcq_prompt = PromptTemplate(
    input_variables=["context"],
    template="""
Create 5 multiple-choice questions from the context below.
Each question has exactly four options. Mark the correct option by enclosing its letter in (**bold**).

Context:
{context}
""",
)

sa_prompt = PromptTemplate(
    input_variables=["context"],
    template="""
Generate 5 short-answer questions answerable directly from the context. Only write questions in your answers and nothing else.

Context:
{context}
""",
)

curiosity_prompt = PromptTemplate(
    input_variables=["text"],
    template="""
Context: {text}

Write EXACTLY five curiosity-driven questions that:
• Are NOT directly answered in the text.
• Connect the text to broader implications, applications or consequences.
• Require higher-order thinking.

Return ONLY the questions, one per line – no numbering or extra text.
""",
)


def _pdf_loader(path: str) -> List[Document]:
    return PyMuPDFLoader(path).load()


def _docx_loader(path: str) -> List[Document]:
    # Strategy 1 – LangChain loader which relies on `docx2txt`.
    try:
        return Docx2txtLoader(path).load()
    except ModuleNotFoundError:
        pass  # fall through – docx2txt missing
    except Exception:
        pass  # any other error, try fallback

    # Strategy 2 – python-docx
    try:
        from docx import Document as _Docx  # type: ignore

        text = "\n".join(p.text for p in _Docx(path).paragraphs)
        return [Document(page_content=text, metadata={"source": path})]
    except ModuleNotFoundError:
        pass
    except Exception:
        pass

    # Strategy 3 – minimal XML extraction (no external deps)
    try:
        import zipfile
        import xml.etree.ElementTree as ET

        with zipfile.ZipFile(path) as zf:
            xml_content = zf.read("word/document.xml")
        tree = ET.fromstring(xml_content)
        text = "\n".join(t.text for t in tree.iter() if t.text)
        return [Document(page_content=text, metadata={"source": path})]
    except Exception as e:
        raise RuntimeError(f"DOCX parse failed: {e}") from e


def _smart_load(path: str) -> List[Document]:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return _pdf_loader(path)
    if ext == ".docx":
        return _docx_loader(path)
    raise ValueError("Unsupported extension")


# NLP helper functions


def _clean(line: str) -> str:
    return re.sub(r"^[\s\-\d\.\)\(•]+", "", line).strip()


def _rank_by_novelty(qs: List[str], q_emb: np.ndarray, d_emb: np.ndarray) -> List[str]:
    """Rank questions by how dissimilar they are to the document embedding."""
    q_emb /= np.linalg.norm(q_emb, axis=1, keepdims=True) + 1e-10
    d_emb /= np.linalg.norm(d_emb, axis=1, keepdims=True) + 1e-10
    novelty = 1 - (q_emb @ d_emb.T).max(axis=1)
    return [q for _, q in sorted(zip(novelty, qs), reverse=True)]

#Core pipeline

def process_file(path: str) -> Dict[str, str]:
    docs = _smart_load(path)
    whole_text = "\n".join(d.page_content for d in docs)

    summary_chain = load_summarize_chain(llm, chain_type="map_reduce")
    summary = summary_chain.run(docs)
    mcqs = LLMChain(llm=llm, prompt=mcq_prompt).run(context=whole_text)
    saqs = LLMChain(llm=llm, prompt=sa_prompt).run(context=whole_text)

    doc_emb = embedder.embed_documents([d.page_content for d in docs])
    raw_curiosity = LLMChain(llm=llm, prompt=curiosity_prompt).run(text=whole_text)

    questions = [_clean(l) for l in raw_curiosity.splitlines() if _clean(l).endswith("?")]
    if questions:
        q_emb = embedder.embed_documents(questions)
        questions = _rank_by_novelty(
            questions,
            np.asarray(q_emb, dtype=np.float32),
            np.asarray(doc_emb, dtype=np.float32),
        )
        curiosity = "\n".join(questions)
    else:
        curiosity = raw_curiosity

    return {
        "summary": summary,
        "mcqs": mcqs,
        "questions": saqs,
        "curiosity": curiosity,
    }

app = Flask(__name__)
app.secret_key = "change-me"
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024  # 25 MB

ALLOWED_EXTENSIONS = {".pdf", ".docx"}


@app.route("/", methods=["GET", "POST"])
def index():
    output = {"summary": "", "mcqs": "", "questions": "", "curiosity": ""}

    if request.method == "POST":
        temp_path = None
        try:
            uploaded = request.files.get("pdf")
            ext = Path(uploaded.filename or "").suffix.lower()
            if not uploaded or ext not in ALLOWED_EXTENSIONS:
                flash("Upload a PDF or DOCX file.", "danger")
                return render_template("index.html", **output)

            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                uploaded.save(tmp.name)
                temp_path = tmp.name

            output = process_file(temp_path)
        except Exception as exc:
            traceback.print_exc()
            flash(f"Processing failed: {exc}", "danger")
        finally:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)

    return render_template("index.html", **output)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", 5000)), debug=False)
